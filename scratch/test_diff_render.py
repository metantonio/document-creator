import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def test_diff_docx():
    doc = Document()
    doc.add_heading("Git Diff Code Block Rendering Test", level=1)
    
    code_lines = [
        "diff --git a/doc_service.py b/doc_service.py",
        "index a1b2c3d..e4f5g6h 100644",
        "--- a/doc_service.py",
        "+++ b/doc_service.py",
        "@@ -10,4 +10,6 @@ def parse_docx(filepath: str):",
        "     doc = Document(filepath)",
        "-    # Old monochrome code block",
        "+    # Enhanced Git Diff code block with green/red line highlighting",
        "+    bg_color = 'DCFCE7' if line.startswith('+') else 'FEE2E2'",
        "     return doc"
    ]
    
    for code_line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        
        run = p.add_run(code_line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        
        bg_color = "F8FAFC"
        text_rgb = RGBColor(51, 65, 85)
        
        if code_line.startswith('+') and not code_line.startswith('+++'):
            bg_color = "DCFCE7" # Soft Light Green
            text_rgb = RGBColor(22, 101, 52) # Dark Green #166534
        elif code_line.startswith('-') and not code_line.startswith('---'):
            bg_color = "FEE2E2" # Soft Light Red
            text_rgb = RGBColor(153, 27, 27) # Dark Red #991B1B
        elif code_line.startswith('@@') or code_line.startswith('diff --git') or code_line.startswith('index '):
            bg_color = "E0E7FF" # Soft Indigo
            text_rgb = RGBColor(55, 48, 163) # Dark Indigo #3730A3
            run.font.bold = True
        elif code_line.startswith('+++') or code_line.startswith('---'):
            bg_color = "F1F5F9"
            text_rgb = RGBColor(71, 85, 105)
            run.font.bold = True

        run.font.color.rgb = text_rgb
        try:
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            p._p.get_or_add_pPr().append(shd)
        except Exception:
            pass

    doc.save("docs_storage/test_diff_render.docx")
    print("Saved test_diff_render.docx successfully!")

if __name__ == "__main__":
    test_diff_docx()
