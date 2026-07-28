import sys
import docx
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

def inspect_doc():
    doc = Document("docs_storage/GPC Technical.docx")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")

    styles_count = {}
    heading_candidates = []

    for idx, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else "None"
        styles_count[style] = styles_count.get(style, 0) + 1

        text = p.text.strip()
        if text.startswith("#") or "Heading" in style or "Title" in style or (len(text) < 80 and p.runs and any(r.bold for r in p.runs) and not text.endswith(".")):
            heading_candidates.append((idx, style, text[:80]))

    print("\nStyles Breakdown:")
    for style, count in styles_count.items():
        print(f"  - {style}: {count}")

    print(f"\nHeading Candidates found: {len(heading_candidates)}")
    for idx, style, text in heading_candidates[:30]:
        print(f"  Line {idx} [{style}]: {text}")

if __name__ == "__main__":
    inspect_doc()
