import os
import re
from typing import List, Dict, Any, Optional
import docx
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor

def normalize_path(path: str) -> str:
    return os.path.abspath(path)

def list_all_documents(document_paths: List[str]) -> List[Dict[str, Any]]:
    """Scan configured document paths for .md, .txt, and .docx files."""
    documents = []
    seen_paths = set()

    for dir_path in document_paths:
        abs_dir = normalize_path(dir_path)
        if not os.path.exists(abs_dir):
            continue
        
        for root, _, files in os.walk(abs_dir):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in ['.md', '.docx', '.txt']:
                    full_path = os.path.join(root, file)
                    if full_path in seen_paths:
                        continue
                    seen_paths.add(full_path)
                    
                    try:
                        stat = os.stat(full_path)
                        documents.append({
                            "filename": file,
                            "filepath": full_path,
                            "folder": root,
                            "format": ext[1:], # 'md', 'docx', 'txt'
                            "size_bytes": stat.st_size,
                            "modified_time": stat.st_mtime
                        })
                    except Exception as e:
                        print(f"Error reading file stat for {full_path}: {e}")

    documents.sort(key=lambda x: x["modified_time"], reverse=True)
    return documents


def parse_markdown_or_txt(content: str) -> List[Dict[str, Any]]:
    """Parse headings and sections from Markdown or Text content."""
    lines = content.splitlines()
    sections = []
    current_title = "Document Overview"
    current_level = 1
    current_lines = []

    for line in lines:
        match = re.match(r'^(#{1,6})\s+(.+)$', line.strip())
        if match:
            if current_lines or sections:
                sections.append({
                    "title": current_title,
                    "level": current_level,
                    "content": "\n".join(current_lines).strip()
                })
            current_level = len(match.group(1))
            current_title = match.group(2).strip()
            current_lines = []
        else:
            txt_heading_match = re.match(r'^([A-Z0-9\s_\-]{3,50}):$', line.strip())
            if txt_heading_match and not line.startswith('#'):
                if current_lines or sections:
                    sections.append({
                        "title": current_title,
                        "level": current_level,
                        "content": "\n".join(current_lines).strip()
                    })
                current_level = 2
                current_title = txt_heading_match.group(1).strip()
                current_lines = []
            else:
                current_lines.append(line)

    if current_lines or not sections:
        sections.append({
            "title": current_title,
            "level": current_level,
            "content": "\n".join(current_lines).strip()
        })

    return sections


def parse_docx(filepath: str) -> List[Dict[str, Any]]:
    """Parse headings, section text, and tables in exact document order from a .docx file."""
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    doc = Document(filepath)
    sections = []
    current_title = "Document Overview"
    current_level = 1
    current_lines = []

    for child in doc.element.body:
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            style_name = p.style.name if p.style else ""
            text = p.text.strip()
            
            is_md_heading = None
            if text.startswith('#'):
                m_h = re.match(r'^(#{1,6})\s+([A-Za-z0-9][^\-\=\#].*)$', text)
                if m_h and not m_h.group(2).strip().startswith('--'):
                    is_md_heading = m_h

            is_bold_title = (len(text) < 80 and p.runs and all(r.bold for r in p.runs) and not text.endswith('.') and not text.startswith('-') and not text.startswith('#'))

            if style_name.startswith('Heading') or style_name == 'Title' or is_md_heading or is_bold_title:
                if current_lines or sections:
                    sections.append({
                        "title": current_title,
                        "level": current_level,
                        "content": "\n".join(current_lines).strip()
                    })
                
                if is_md_heading:
                    current_level = len(is_md_heading.group(1))
                    current_title = is_md_heading.group(2).strip()
                elif style_name == 'Title' or (is_bold_title and current_level == 1):
                    current_level = 1
                    current_title = text if text else "Untitled Section"
                else:
                    try:
                        current_level = int(style_name.replace('Heading', '').strip())
                    except ValueError:
                        current_level = 2
                    current_title = text if text else "Untitled Section"
                current_lines = []
            else:
                if p.text:
                    current_lines.append(p.text)
                    
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            table_lines = []
            for r_idx, row in enumerate(table.rows):
                row_vals = [c.text.strip().replace('\n', ' ') for c in row.cells]
                table_lines.append("| " + " | ".join(row_vals) + " |")
                if r_idx == 0:
                    table_lines.append("| " + " | ".join(["---"] * len(row_vals)) + " |")
            if table_lines:
                current_lines.append("\n" + "\n".join(table_lines) + "\n")

    if current_lines or not sections:
        sections.append({
            "title": current_title,
            "level": current_level,
            "content": "\n".join(current_lines).strip()
        })

    return sections


def read_document(filepath: str) -> Dict[str, Any]:
    """Read document content and extract structural headings/sections."""
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.docx':
        doc = Document(filepath)
        full_text_parts = [p.text for p in doc.paragraphs if p.text]
        for t in doc.tables:
            for r in t.rows:
                full_text_parts.append(" | ".join([c.text.strip() for c in r.cells]))
        full_text = "\n\n".join(full_text_parts)
        sections = parse_docx(filepath)
    else:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            full_text = f.read()
        sections = parse_markdown_or_txt(full_text)

    headings = [{"title": s["title"], "level": s["level"]} for s in sections]

    return {
        "filepath": filepath,
        "filename": os.path.basename(filepath),
        "format": ext[1:],
        "full_text": full_text,
        "headings": headings,
        "sections": sections
    }


def render_inline_markdown(paragraph, text: str):
    """Parse inline bold (**text**), italic (*text*), code (`code`), and link ([text](url)) into docx runs."""
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(71, 85, 105)
        elif token.startswith('[') and ']' in token and '(' in token and token.endswith(')'):
            link_text = token[1:token.index(']')]
            url = token[token.index('(')+1:-1]
            run = paragraph.add_run(f"{link_text} ({url})")
            run.font.color.rgb = RGBColor(5, 99, 193)
            run.underline = True
        else:
            paragraph.add_run(token)


def render_markdown_body_to_docx(doc: Document, content: str):
    """
    Renders markdown section body into native Word elements:
    - Tables -> Native docx tables with grid borders, headers, and alternating fills
    - Code Blocks -> Native code blocks with Consolas font and background shading
    - Bullet / Numbered Lists -> List Bullet / List Number styles
    - Paragraphs -> Formatted runs with bold/italic/links
    """
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        if not line:
            i += 1
            continue
            
        # 1. Fenced Code Blocks (```python / ```diff / ```)
        if line.startswith('```'):
            code_lang = line.replace('```', '').strip().lower()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines) and lines[i].startswith('```'):
                i += 1

            # Detect Git Diff format
            is_diff = (code_lang.startswith('diff')) or any(
                cl.startswith('+') or cl.startswith('-') or cl.startswith('@@') or cl.startswith('diff --git')
                for cl in code_lines
            )

            if is_diff:
                for code_line in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(12)
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    
                    run = p.add_run(code_line if code_line else ' ')
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)

                    bg_color = "F8FAFC"
                    text_rgb = RGBColor(51, 65, 85)

                    if code_line.startswith('+') and not code_line.startswith('+++'):
                        bg_color = "DCFCE7"  # Soft Light Green
                        text_rgb = RGBColor(22, 101, 52)  # Dark Green #166534
                    elif code_line.startswith('-') and not code_line.startswith('---'):
                        bg_color = "FEE2E2"  # Soft Light Red
                        text_rgb = RGBColor(153, 27, 27)  # Dark Red #991B1B
                    elif code_line.startswith('@@') or code_line.startswith('diff --git') or code_line.startswith('index '):
                        bg_color = "E0E7FF"  # Soft Indigo
                        text_rgb = RGBColor(55, 48, 163)  # Dark Indigo #3730A3
                        run.font.bold = True
                    elif code_line.startswith('+++') or code_line.startswith('---'):
                        bg_color = "F1F5F9"
                        text_rgb = RGBColor(71, 85, 105)
                        run.font.bold = True

                    run.font.color.rgb = text_rgb
                    try:
                        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
                        p._p.get_or_add_pPr().append(shd)
                    except Exception:
                        pass
            else:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(30, 41, 59)
                
                try:
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                    p._p.get_or_add_pPr().append(shd)
                except Exception:
                    pass
            continue

        # 1.5 Callouts / Blockquotes (> [!NOTE] text...)
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(re.sub(r'^\s*>\s?', '', lines[i]))
                i += 1
            
            quote_text = "\n".join(quote_lines).strip()
            
            border_color = "4F46E5" # Indigo Note
            bg_color = "EEF2FF"
            callout_title = "NOTE"
            
            if quote_text.startswith('[!WARNING]') or quote_text.startswith('[!CAUTION]'):
                border_color = "EF4444" # Red
                bg_color = "FEF2F2"
                callout_title = "WARNING"
                quote_text = re.sub(r'^\[!(WARNING|CAUTION)\]\s*', '', quote_text)
            elif quote_text.startswith('[!IMPORTANT]'):
                border_color = "F59E0B" # Amber
                bg_color = "FFFBEB"
                callout_title = "IMPORTANT"
                quote_text = re.sub(r'^\[!IMPORTANT\]\s*', '', quote_text)
            elif quote_text.startswith('[!TIP]'):
                border_color = "10B981" # Green
                bg_color = "ECFDF5"
                callout_title = "TIP"
                quote_text = re.sub(r'^\[!TIP\]\s*', '', quote_text)
            elif quote_text.startswith('[!NOTE]'):
                quote_text = re.sub(r'^\[!NOTE\]\s*', '', quote_text)

            callout_table = doc.add_table(rows=1, cols=1)
            callout_cell = callout_table.rows[0].cells[0]
            cp = callout_cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after = Pt(2)
            
            run_lbl = cp.add_run(f"📌 {callout_title}: ")
            run_lbl.font.bold = True
            run_lbl.font.color.rgb = RGBColor.from_string(border_color)
            
            render_inline_markdown(cp, quote_text)
            
            try:
                tcPr = callout_cell._tc.get_or_add_tcPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
                borders = parse_xml(f'''
                    <w:tcBorders {nsdecls("w")}>
                        <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
                        <w:top w:val="none"/>
                        <w:right w:val="none"/>
                        <w:bottom w:val="none"/>
                    </w:tcBorders>
                ''')
                tcPr.append(shd)
                tcPr.append(borders)
            except Exception:
                pass
            continue

        # 2. Markdown Tables (| Col1 | Col2 |)
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
                
            parsed_rows = []
            for t_line in table_lines:
                if re.match(r'^\|[\s:\-|\+]+\|$', t_line):
                    continue
                cells = [c.strip() for c in t_line.strip('|').split('|')]
                parsed_rows.append(cells)
                
            if parsed_rows:
                num_cols = max(len(row) for row in parsed_rows)
                num_rows = len(parsed_rows)
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                
                for r_idx, row_cells in enumerate(parsed_rows):
                    row = table.rows[r_idx]
                    
                    # Prevent row split across pages
                    try:
                        trPr = row._tr.get_or_add_trPr()
                        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                    except Exception:
                        pass

                    if r_idx == 0:
                        try:
                            trPr = row._tr.get_or_add_trPr()
                            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
                        except Exception:
                            pass

                    for c_idx, cell_value in enumerate(row_cells):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            render_inline_markdown(p, cell_value)
                            
                            if r_idx == 0: # Header Row
                                for run in p.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                try:
                                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4F46E5"/>')
                                    cell._tc.get_or_add_tcPr().append(shd)
                                except Exception:
                                    pass
                            elif r_idx % 2 == 1: # Alternating row fill
                                try:
                                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
                                    cell._tc.get_or_add_tcPr().append(shd)
                                except Exception:
                                    pass
            continue

        # 3. Headings inside content (#, ##, ###)
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        # 4. Bullet / Numbered Lists
        bullet_match = re.match(r'^\s*[\-\*]\s+(.+)$', line)
        number_match = re.match(r'^\s*\d+\.\s+(.+)$', line)
        
        if bullet_match:
            p = doc.add_paragraph(style='List Bullet')
            render_inline_markdown(p, bullet_match.group(1).strip())
            i += 1
            continue
        elif number_match:
            p = doc.add_paragraph(style='List Number')
            render_inline_markdown(p, number_match.group(1).strip())
            i += 1
            continue

        # 5. Standard Paragraphs
        p = doc.add_paragraph()
        render_inline_markdown(p, line)
        i += 1


def create_document(filepath: str, format_type: str, title: str, content: str) -> Dict[str, Any]:
    """Create a new document file (.md, .txt, or .docx)."""
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    format_type = format_type.lower()

    if format_type == "docx":
        doc = Document()
        doc.add_heading(title, level=1)
        render_markdown_body_to_docx(doc, content)
        doc.save(filepath)
    elif format_type == "md":
        formatted_content = f"# {title}\n\n{content}"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(formatted_content)
    else: # txt
        formatted_content = f"{title.upper()}\n{'='*len(title)}\n\n{content}"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(formatted_content)

    return read_document(filepath)


def save_updated_sections(filepath: str, format_type: str, sections: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Save an updated list of sections back to disk in specified format."""
    format_type = format_type.lower()
    
    if format_type == "docx":
        doc = Document()
        for sec in sections:
            level = max(1, min(sec.get("level", 2), 4))
            doc.add_heading(sec["title"], level=level)
            body = sec.get("content", "")
            if body:
                render_markdown_body_to_docx(doc, body)
        doc.save(filepath)
    elif format_type == "md":
        md_text = ""
        for sec in sections:
            prefix = "#" * max(1, min(sec.get("level", 2), 6))
            md_text += f"{prefix} {sec['title']}\n\n{sec.get('content', '')}\n\n"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(md_text.strip() + "\n")
    else: # txt
        txt_text = ""
        for sec in sections:
            txt_text += f"{sec['title'].upper()}:\n{'-'*len(sec['title'])}\n{sec.get('content', '')}\n\n"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(txt_text.strip() + "\n")

    return read_document(filepath)
